from flask import Flask, render_template, request, redirect, url_for, send_file, flash, session, send_from_directory, jsonify, make_response
from sqlalchemy.orm import sessionmaker
from sqlalchemy import String, func, text
import pandas as pd
from io import BytesIO
import io
from datetime import datetime, date, timedelta
from werkzeug.security import check_password_hash, generate_password_hash
import logging
from contextlib import contextmanager
from database import get_engine
from models import Contrato, usuarios, Spei, Retiros, Nominales
import os
from pathlib import Path
from crud import clean_sql_string
import random, string
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)
app.secret_key = 'your_secret_key'
app.permanent_session_lifetime = timedelta(minutes=15)
UPLOAD_FOLDER = "temp_uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Configuración de la conexión a la base de datos
engine = get_engine()
Session = sessionmaker(bind=engine)
logging.basicConfig(level=logging.INFO)

@contextmanager
def get_session():
    session = Session()
    try:
        yield session
    finally:
        session.close()

@app.route("/register", methods=["GET", "POST"])
def register():
    if "usuario" in session:
        return redirect(url_for("index"))

    if request.method == "POST":
        nombre = request.form["nombre"]
        correo = request.form["correo"]
        contrasena = request.form["contraseña"]
        hashed_password = generate_password_hash(contrasena)

        with get_session() as session_db:
            if session_db.query(usuarios).filter_by(correo=correo).first():
                flash("El usuario ya existe", "error")
                return redirect("/register")

            new_user = usuarios(nombre=nombre, correo=correo, contrasena=hashed_password)
            session_db.add(new_user)
            session_db.commit()
            flash("Usuario registrado con éxito", "success")
            return redirect("/login")

    return render_template("register.html")

@app.route("/descargar")
def descargar_archivo():
    return send_from_directory("static", "SOLICITUD DE CONTRATOS.xlsx", as_attachment=True)

@app.route("/login", methods=["GET", "POST"])
def login():
    if "usuario" in session:
        return redirect(url_for("index"))

    if request.method == "POST":
        correo = request.form["correo"]
        contrasena = request.form["contrasena"]

        with get_session() as session_db:
            user = session_db.query(usuarios).filter_by(correo=correo).first()

            if user and check_password_hash(user.contrasena, contrasena):
                session["usuario"] = user.nombre
                session.permanent = True
                flash("Inicio de sesión exitoso", "success")
                return redirect(url_for("index"))

            flash("Correo o contraseña incorrectos.", "error")
            return redirect('/login')

    return render_template("login.html")

@app.route("/", methods=["GET", "POST"])
def index():
    if "usuario" not in session:
        return redirect("/login")
    
    return render_template(
        "index.html",
    )

@app.route("/concentradocontratos", methods=["GET", "POST"])
def concentradocontratos():
    if "usuario" not in session:
        return redirect("/login")

    usuario_logeado = session["usuario"]
    search_term = request.args.get("search", "")
    page = int(request.args.get("page", 1))
    per_page = 15

    with get_session() as session_db:
        # Verificar si el usuario es maestro
        user = session_db.query(usuarios).filter_by(nombre=usuario_logeado).first()
        es_maestro = user and user.es_maestro

        # Modificar la consulta según el tipo de usuario
        if es_maestro:
            query = session_db.query(Contrato)
        else:
            query = session_db.query(Contrato).filter_by(usuario=usuario_logeado)

        # Filtro de búsqueda
        if search_term:
            query = query.filter(
                Contrato.numero_cuenta.cast(String).like(f"%{search_term}%") |
                Contrato.titular.like(f"%{search_term}%") |
                Contrato.motivo.like(f"%{search_term}%") |
                Contrato.estatus.like(f"%{search_term}%") |
                Contrato.sucursal.like(f"%{search_term}%") |
                Contrato.usuario.like(f"%{search_term}%")
            )

        # Paginación y obtención de datos
        total = query.count()
        contratos = query.order_by(Contrato.fecha_carga.desc()).offset((page - 1) * per_page).limit(per_page).all()

        contratos_data = [
            {
                "id": contrato.id,
                "numero_remesa": contrato.numero_remesa or '',
                "numero_tarjeta": contrato.numero_tarjeta or '',
                "numero_cuenta": contrato.numero_cuenta or '',
                "titular": contrato.titular or '',
                "nivel_cuenta": contrato.nivel_cuenta or '',
                "cc": contrato.cc or '',
                "sucursal": contrato.sucursal or '',
                "motivo": contrato.motivo or '',
                "observaciones": contrato.observaciones or '',
                "fecha_carga": contrato.fecha_carga or '',
                "estatus": contrato.estatus or '',
                "url": contrato.url or '',
                "usuario": contrato.usuario or ''
            }
            for contrato in contratos
        ]

        # Obtener el conteo de estatus para usuarios maestros o regulares
        if es_maestro:
            estatus_count = session_db.query(
                Contrato.estatus, func.count(Contrato.id).label("total")
            ).group_by(Contrato.estatus).order_by(Contrato.estatus).all()
        else:
            estatus_count = session_db.query(
                Contrato.estatus, func.count(Contrato.id).label("total")
            ).filter_by(usuario=usuario_logeado).group_by(Contrato.estatus).order_by(Contrato.estatus).all()

        # Convertir los resultados a un diccionario de totales por estatus
        estatus_totals = {row.estatus: row.total for row in estatus_count}

        # Calcular el total general a partir de los resultados de estatus_count
        total_general = sum(estatus_totals.values())

        # Agregar el total general al diccionario para facilitar su acceso
        estatus_totals["total_general"] = total_general

    return render_template(
        "concentradocontratos.html",
        contratos=contratos_data,
        page=page,
        total=total,
        per_page=per_page,
        search_term=search_term,
        estatus_totals=estatus_totals,
        total_general=total_general,
        usuario_logeado=session["usuario"],
    )

@app.route("/logout", methods=["POST"])
def logout():
    session.clear()
    flash("Has cerrado sesión exitosamente.", "success")
    return redirect('/login')

@app.before_request
def make_session_permanent():
    if request.endpoint not in ['login', 'register', 'static']:
        if "usuario" not in session:
            flash("Tu sesión ha expirado. Por favor, inicia sesión nuevamente.", "error")
            return redirect("/login")
        session.permanent = True

@app.route('/upload', methods=['GET', 'POST'])
def upload_file():
    if "usuario" not in session:
        return redirect("/login")

    if request.method == 'GET':
        return render_template('upload.html')

    try:
        if 'file' not in request.files:
            flash('No se seleccionó ningún archivo.', 'error')
            return redirect('/upload')

        file = request.files['file']
        if file.filename == '':
            flash('El nombre del archivo está vacío.', 'error')
            return redirect(url_for('upload'))

        df = pd.read_excel(file)
        required_columns = ["NUMERO DE CUENTA A 10 DIGITOS", "TITULAR", "CC"]
        for column in required_columns:
            if column not in df.columns:
                flash(f"El archivo no contiene la columna obligatoria: {column}", 'error')
                return redirect('/upload')

        # Limpiar caracteres en columnas de texto
        text_columns = ["NUMERO DE REMESA", "TITULAR", "SUCURSAL", "MOTIVO", "OBSERVACIONES"]
        for column in text_columns:
            if column in df.columns:
                df[column] = df[column].apply(clean_sql_string)

        df = df.dropna(subset=required_columns)
        if df.empty:
            flash("No hay filas válidas en el archivo después de filtrar valores nulos.", 'error')
            return redirect('/upload')

        # Reemplazar valores NaN con cadenas vacías en todo el DataFrame
        df = df.fillna("")

        numeric_columns = [
            "NUMERO DE CUENTA A 10 DIGITOS",
            "NIVEL DE CUENTA",
            "CC",
            "NUMERO DE TARJETA",
        ]
        for column in numeric_columns:
            if column in df.columns:
                # Primero limpiar cualquier carácter no numérico
                df[column] = df[column].astype(str).apply(lambda x: ''.join(filter(str.isdigit, str(x))))
                df[column] = pd.to_numeric(df[column], errors='coerce').fillna(0).astype(int)

        estatus = "SOLICITADO"
        fecha_carga = date.today()
        duplicados = []
        errores_cc = []
        invalid_rows = []
        caracteres_especiales = []

        with get_session() as session_db:
            for _, row in df.iterrows():
                sucursal = clean_sql_string(row["SUCURSAL"])
                numero_cuenta = row["NUMERO DE CUENTA A 10 DIGITOS"]
                cc = row["CC"]

                # Verificar si hubo limpieza de caracteres especiales
                original_titular = str(row["TITULAR"])
                cleaned_titular = clean_sql_string(original_titular)
                if original_titular != cleaned_titular:
                    caracteres_especiales.append(f"Cuenta: {numero_cuenta}")

                if len(str(numero_cuenta)) < 8 or len(str(numero_cuenta)) > 10:
                    invalid_rows.append(row.to_dict())
                    continue

                if cc == 0:
                    errores_cc.append(f"Cuenta: {numero_cuenta} con CC igual a 0.")
                    continue

                if session_db.query(Contrato).filter_by(numero_cuenta=numero_cuenta).first():
                    duplicados.append(numero_cuenta)
                    continue

                if pd.isna(sucursal) or sucursal == "":
                    sucursal = "SUCURSAL NO CARGADA CORRECTAMENTE"
                    continue

                contrato = Contrato(
                    numero_remesa=row.get("NUMERO DE REMESA"),
                    numero_tarjeta=row.get("NUMERO DE TARJETA"),
                    numero_cuenta=numero_cuenta,
                    titular=cleaned_titular,
                    nivel_cuenta=row.get("NIVEL DE CUENTA", 0),
                    cc=cc,
                    sucursal=sucursal,
                    motivo=clean_sql_string(row.get("MOTIVO", "")),
                    observaciones=clean_sql_string(row.get("OBSERVACIONES", "")),
                    fecha_carga=fecha_carga,
                    usuario=session["usuario"],
                    estatus=estatus
                )
                session_db.add(contrato)

            session_db.commit()

        if caracteres_especiales:
            flash(f"Se limpiaron caracteres especiales en {len(caracteres_especiales)} registros", 'info')

        if duplicados:
            flash(f"Algunas cuentas ya existían y no se cargaron: {', '.join(map(str, duplicados))}", 'warning')

        if errores_cc:
            flash(f"Errores de CC (igual a 0): {', '.join(errores_cc)}", 'error')

        if invalid_rows:
            flash(f"Filas inválidas (número de cuenta no tiene entre 8 y 10 dígitos): {len(invalid_rows)}", 'warning')

        flash("Archivo procesado y datos cargados en la base de datos con éxito.", 'success')
    except Exception as e:
        flash(f"Error al procesar el archivo: {str(e)}", 'error')

    return redirect('/upload')

@app.route("/update-status/<int:contrato_id>", methods=["POST"])
def update_status(contrato_id):
    if "usuario" not in session:
        return redirect("/login")

    nuevo_estatus = request.form.get("estatus")
    if not nuevo_estatus:
        flash("El estatus es obligatorio.", "error")
        return redirect(request.referrer)

    with get_session() as session_db:
        contrato = session_db.query(Contrato).filter_by(id=contrato_id).first()
        if contrato:
            contrato.estatus = nuevo_estatus
            session_db.commit()
            flash(f"Contrato {contrato_id} actualizado a {nuevo_estatus}.", "success")
        else:
            flash("Contrato no encontrado.", "error")

    return redirect("/concentradocontratos")

@app.route("/update", methods=["GET", "POST"])
def update_massive_status():
    if "usuario" not in session:
        return redirect("/login")
    
    if request.method == "POST":   
        file = request.files.get("file")
        if not file or file.filename == "":
            flash("Debes cargar un archivo válido.", "error")
            return render_template("update.html")

        try:
            df = pd.read_excel(file)
            column_mapping = {
                "CONTRATO": "numero_cuenta",
                "ESTATUS": "estatus",
                "UBICACIÓN ARCHIVO": "url"
            }
            for col in column_mapping.keys():
                if col not in df.columns:
                    flash(f"El archivo debe contener la columna '{col}'.", "error")
                    return render_template("update.html")

            df.rename(columns=column_mapping, inplace=True)

            df["numero_cuenta"] = pd.to_numeric(df["numero_cuenta"], errors="coerce").fillna(0).astype(int)
            df["estatus"] = df["estatus"].astype(str).str.strip()
            df["url"] = df["url"].astype(str).str.strip()

            df["url"] = df["url"].replace("nan", "").fillna("")

            allowed_statuses = {"GENERADO", "ENVIADO", "OTRO PROGRAMA", "SOLICITADO", "CONTRATO INEXISTENTE"}
            invalid_status = df[~df["estatus"].isin(allowed_statuses)]

            if not invalid_status.empty:
                # Generar archivo con los contratos rechazados
                invalid_details = invalid_status[["numero_cuenta", "estatus", "url"]]
                file_buffer = io.BytesIO()
                invalid_details.to_csv(file_buffer, index=False)
                file_buffer.seek(0)

                flash(
                    "Algunos contratos tienen un estatus inválido. Se generó un archivo con los detalles.",
                    "error"
                )
                return send_file(
                    file_buffer,
                    as_attachment=True,
                    download_name="estatus_invalidos.csv",
                    mimetype="text/csv",
                )

            missing_cuentas = []
            updated_count = 0
            invalid_url = []

            with get_session() as session_db:
                for _, row in df.iterrows():
                    numero_cuenta = row["numero_cuenta"]
                    estatus = row["estatus"]
                    url = row["url"]

                    contrato = session_db.query(Contrato).filter_by(numero_cuenta=numero_cuenta).first()
                    if not contrato:
                        missing_cuentas.append(numero_cuenta)
                        continue
                    
                    # Validaciones según el estatus
                    if estatus in {"ENVIADO", "SOLICITADO"}:
                        if not url:
                            url = contrato.url
                    elif estatus == "GENERADO" and not url:
                        invalid_url.append(numero_cuenta)
                        continue

                    # Actualizar contrato
                    contrato.estatus = estatus
                    contrato.url = url
                    updated_count += 1

                # Commit de la transacción
                if updated_count > 0:
                    session_db.commit()

            # Manejar resultados
            if updated_count > 0:
                flash(f"Estatus actualizado para {updated_count} contrato(s).", "success")
            
            if invalid_url:
                flash(f"Algunos contratos con estatus 'GENERADO' no tienen una ubicación válida. Revisa el archivo.", "warning")
                file_buffer = io.BytesIO()
                pd.DataFrame(invalid_url, columns=["numero_cuenta"]).to_csv(file_buffer, index=False)
                file_buffer.seek(0)
                return send_file(
                    file_buffer,
                    as_attachment=True,
                    download_name="ubicaciones_invalidas.csv",
                    mimetype="text/csv"
                )

            if missing_cuentas:
                cuentas_str = "\n".join(map(str, missing_cuentas))
                file_buffer = io.BytesIO()
                file_buffer.write(cuentas_str.encode('utf-8'))
                file_buffer.seek(0)
                flash("Algunas cuentas no se encontraron en la base de datos.", "warning")
                return send_file(
                    file_buffer,
                    as_attachment=True,
                    download_name="cuentas_no_encontradas.txt",
                    mimetype="text/plain",
                )

        except Exception as e:
            flash(f"Error al procesar el archivo: {e}", "error")
    
    return render_template("update.html")

@app.route('/download', methods=['GET', 'POST'])
def download():
    if "usuario" not in session:
        flash("Por favor, inicia sesión para acceder a esta página.", "error")
        return redirect(url_for("login"))

    usuario_logeado = session["usuario"]

    with get_session() as session_db:
        user = session_db.query(usuarios).filter_by(nombre=usuario_logeado).first()
        es_maestro = user and user.es_maestro

        if request.method == 'POST':
            try:
                start_date = request.form.get('start_date')
                end_date = request.form.get('end_date')
                opcion_descarga = request.form.get('opcion_descarga')

                if not start_date or not end_date:
                    flash("Por favor, selecciona un rango de fechas válido.", "error")
                    return redirect(url_for("download"))

                start_date = datetime.strptime(start_date, '%Y-%m-%d').date()
                end_date = datetime.strptime(end_date, '%Y-%m-%d').date()

                query = session_db.query(Contrato).filter(
                    Contrato.fecha_carga >= start_date, Contrato.fecha_carga <= end_date
                )

                if not es_maestro or (es_maestro and opcion_descarga == "propios"):
                    query = query.filter(Contrato.usuario == usuario_logeado)

                contratos = query.all()

                if not contratos:
                    flash("No se encontraron registros para la consulta.", "info")
                    return redirect(url_for("download"))

                data = [
                    {
                        "NUMERO DE REMESA": contrato.numero_remesa,
                        "NUMERO DE TARJETA": contrato.numero_tarjeta,
                        "NUMERO DE CUENTA": contrato.numero_cuenta,
                        "TITULAR": contrato.titular,
                        "NIVEL DE CUENTA": contrato.nivel_cuenta,
                        "CC": contrato.cc,
                        "SUCURSAL": contrato.sucursal,
                        "MOTIVO": contrato.motivo,
                        "OBSERVACIONES": contrato.observaciones,
                        "FECHA DE CARGA": contrato.fecha_carga,
                        "ESTATUS": contrato.estatus,
                        "UBICACIÓN ARCHIVO": contrato.url,
                        "USUARIO": contrato.usuario
                    }
                    for contrato in contratos
                ]

                df = pd.DataFrame(data)
                output = io.BytesIO()
                with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                    df.to_excel(writer, index=False, sheet_name='Contratos')

                output.seek(0)

                return send_file(
                    output,
                    as_attachment=True,
                    download_name="Concentrado_Contratos.xlsx",
                    mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )

            except Exception as e:
                flash(f"Error al procesar la solicitud: {e}", "error")
                return redirect(url_for("download"))

    return render_template("download.html")

# Ruta para mostrar el formulario
@app.route('/spei', methods=['GET'])
def spei_form():
    if "usuario" not in session:
        return redirect("/login")
    return render_template("spei.html")

# Ruta para manejar el envío del formulario
@app.route('/insert_spei', methods=['POST'])
def insert_spei():
    if "usuario" not in session:
        return redirect("/login")
    try:
        nuevo_spei = Spei(
            fecha_tramite=request.form['fecha_tramite'],
            cc=request.form['cc'],
            sucursal=request.form['sucursal'],
            cuenta_origen=request.form['cuenta_origen'],
            titular=request.form['titular'],
            cuenta_destino=request.form['cuenta_beneficiaria'],
            beneficiario=request.form['beneficiario'],
            importe=request.form['importe'],
            autorizador=request.form['autorizador'],
            usuario=session["usuario"]
        )
        with get_session() as session_db:
            session_db.add(nuevo_spei)
            session_db.commit()
        flash("Registro insertado con éxito", "success")
    except Exception as e:
        flash(f"Error al insertar registro: {e}", "danger")
    return redirect("/spei")

@app.route('/buscar_sucursal', methods=['GET'])
def buscar_sucursal():
    cc = request.args.get('cc', type=int)

    if not cc:
        return jsonify({"message": "El parámetro CC es requerido"}), 400

    engine = get_engine()
    # Consulta SQL con JOIN para obtener la sucursal y el CR
    query = """
        SELECT suc.sucursal, coord.cr
        FROM sucursal suc
        JOIN coordinador coord ON suc.id_coordinador = coord.id_coordinador
        WHERE suc.cc = :cc
    """

    try:
        with engine.connect() as conn:
            result = conn.execute(text(query), {"cc": cc}).fetchone()
            if result:
                logging.info(f"Resultado de la consulta para CC {cc}: {result}")
                return jsonify({"sucursal": result[0], "cr": result[1]})
            else:
                return jsonify({"message": "Sucursal no encontrada"}), 404
    except Exception as e:
        logging.error(f"Error en la consulta para CC {cc}: {e}")
        return jsonify({"message": "Error interno del servidor"}), 500

@app.route("/concentradospei", methods=["GET", "POST"])
def concentradospei():
    if "usuario" not in session:
        return redirect("/login")

    usuario_logeado = session["usuario"]
    search_term = request.args.get("search", "")
    page = int(request.args.get("page", 1))
    per_page = 15

    with get_session() as session_db:
        # Verificar si el usuario es maestro
        user = session_db.query(usuarios).filter_by(nombre=usuario_logeado).first()
        es_maestro = user and user.es_maestro

        # Modificar la consulta según el tipo de usuario
        if es_maestro:
            query = session_db.query(Spei)
        else:
            query = session_db.query(Spei).filter_by(usuario=usuario_logeado)

        # Filtro de búsqueda
        if search_term:
            query = query.filter(
                Spei.cuenta_origen.cast(String).like(f"%{search_term}%") |
                Spei.titular.like(f"%{search_term}%") |
                Spei.beneficiario.like(f"%{search_term}%") |
                Spei.estatus.like(f"%{search_term}%") |
                Spei.sucursal.like(f"%{search_term}%") |
                Spei.usuario.like(f"%{search_term}%") |
                Spei.cuenta_destino.like(f"%{search_term}%")
            )

        # Paginación y obtención de datos
        total = query.count()
        spei = query.order_by(Spei.fecha_tramite.desc()).offset((page - 1) * per_page).limit(per_page).all()

        spei_data = [
            {
                "id": spei.id,
                "fecha_tramite": spei.fecha_tramite or '',
                "cc": spei.cc or '',
                "sucursal": spei.sucursal or '',
                "cuenta_origen": spei.cuenta_origen or '',
                "titular": spei.titular or '',
                "cuenta_destino": spei.cuenta_destino or '',
                "beneficiario": spei.beneficiario or '',
                "importe": spei.importe or '',
                "autorizador": spei.autorizador or '',
                "estatus": spei.estatus or '',
                "usuario": spei.usuario or ''
            }
            for spei in spei
        ]

        # Obtener el conteo de estatus para usuarios maestros o regulares
        if es_maestro:
            estatus_count = session_db.query(
                Spei.estatus, func.count(Spei.id).label("total")
            ).group_by(Spei.estatus).order_by(Spei.estatus).all()
        else:
            estatus_count = session_db.query(
                Spei.estatus, func.count(Spei.id).label("total")
            ).filter_by(usuario=usuario_logeado).group_by(Spei.estatus).order_by(Spei.estatus).all()

        # Convertir los resultados a un diccionario de totales por estatus
        estatus_totals = {row.estatus: row.total for row in estatus_count}

        # Calcular el total general a partir de los resultados de estatus_count
        total_general = sum(estatus_totals.values())

        # Agregar el total general al diccionario para facilitar su acceso
        estatus_totals["total_general"] = total_general

    return render_template(
        "concentradospei.html",
        spei=spei_data,
        page=page,
        total=total,
        per_page=per_page,
        search_term=search_term,
        estatus_totals=estatus_totals,
        total_general=total_general,
        usuario_logeado=session["usuario"],
    )

@app.route('/downloadspei', methods=['GET', 'POST'])
def downloadspei():
    if "usuario" not in session:
        flash("Por favor, inicia sesión para acceder a esta página.", "error")
        return redirect(url_for("login"))

    usuario_logeado = session["usuario"]

    with get_session() as session_db:
        user = session_db.query(usuarios).filter_by(nombre=usuario_logeado).first()
        es_maestro = user and user.es_maestro

        if request.method == 'POST':
            try:
                start_date = request.form.get('start_date')
                end_date = request.form.get('end_date')
                opcion_descarga = request.form.get('opcion_descarga')

                if not start_date or not end_date:
                    flash("Por favor, selecciona un rango de fechas válido.", "error")
                    return redirect(url_for("downloadspei"))

                start_date = datetime.strptime(start_date, '%Y-%m-%d').date()
                end_date = datetime.strptime(end_date, '%Y-%m-%d').date()

                query = session_db.query(Spei).filter(
                    Spei.fecha_tramite >= start_date, Spei.fecha_tramite <= end_date
                )

                if not es_maestro or (es_maestro and opcion_descarga == "propios"):
                    query = query.filter(Spei.usuario == usuario_logeado)

                spei = query.all()

                if not spei:
                    flash("No se encontraron registros para la consulta.", "info")
                    return redirect(url_for("downloadspei"))

                data = [
                    {
                        "fecha_tramite": spei.fecha_tramite,
                        "cc": spei.cc,
                        "sucursal": spei.sucursal,
                        "cuenta_origen": spei.cuenta_origen,
                        "titular": spei.titular,
                        "cuenta_destino": spei.cuenta_destino,
                        "beneficiario": spei.beneficiario,
                        "importe": spei.importe,
                        "autorizador": spei.autorizador,
                        "estatus": spei.estatus,
                        "usuario": spei.usuario
                    }
                    for spei in spei
                ]

                df = pd.DataFrame(data)
                output = io.BytesIO()
                with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                    df.to_excel(writer, index=False, sheet_name='Speis')

                output.seek(0)

                return send_file(
                    output,
                    as_attachment=True,
                    download_name="Concentrado_Spei.xlsx",
                    mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )

            except Exception as e:
                flash(f"Error al procesar la solicitud: {e}", "error")
                return redirect(url_for("downloadspei"))

    return render_template("downloadspei.html")

@app.route("/statusspei/<int:spei_id>", methods=["POST"])
def update_statusspei(spei_id):
    if "usuario" not in session:
        return redirect("/login")

    nuevo_estatus = request.form.get("estatus")
    if not nuevo_estatus:
        flash("El estatus es obligatorio.", "error")
        return redirect(request.referrer)

    with get_session() as session_db:
        spei = session_db.query(Spei).filter_by(id=spei_id).first()
        if spei:
            spei.estatus = nuevo_estatus
            session_db.commit()
            flash(f"Spei {spei_id} actualizado a {nuevo_estatus}.", "success")
        else:
            flash("Spei no encontrado.", "error")
            return redirect("/concentradospei")

    return redirect("/concentradospei")

@app.route('/retiros', methods=['GET'])
def retiros_form():
    if "usuario" not in session:
        return redirect("/login")
    return render_template("retiros.html")

@app.route('/insert_retiro', methods=['POST'])
def insert_retiro():
    if "usuario" not in session:
        flash("Sesión expirada. Por favor inicie sesión nuevamente.", "danger")
        return redirect("/login")
    
    try:
        # Validar campos requeridos
        required_fields = [
            'fecha_tramite', 'cc', 'sucursal', 'cr', 'cuenta_origen',
            'titular', 'cuenta_beneficiaria', 'beneficiario', 'importe',
            'clave_autorizacion', 'operacion', 'motivo'
        ]
        
        form_data = {field: request.form.get(field) for field in required_fields}
        missing_fields = [field for field, value in form_data.items() if not value]
        if missing_fields:
            flash(f"Los siguientes campos son requeridos: {', '.join(missing_fields)}", "danger")
            return redirect("/retiros")
        
        # Validación de cuentas beneficiarias
        if form_data['cuenta_beneficiaria'] != request.form.get('confirmar_cuenta_beneficiaria'):
            flash("Las cuentas beneficiarias no coinciden.", "danger")
            return redirect("/retiros")
        
        # Validación de importe y fecha
        try:
            form_data['fecha_tramite'] = datetime.strptime(form_data['fecha_tramite'], '%Y-%m-%d')
            form_data['importe'] = float(form_data['importe'])
            if form_data['importe'] <= 0:
                raise ValueError("El importe debe ser mayor a 0")
        except ValueError as e:
            flash(str(e), "danger")
            return redirect("/retiros")
        
        # Crear y guardar el objeto Retiros
        nuevo_retiro = Retiros(
            fecha_tramite=form_data['fecha_tramite'],
            cc=form_data['cc'],
            sucursal=form_data['sucursal'],
            cr=form_data['cr'],
            cuenta_origen=form_data['cuenta_origen'],
            titular=form_data['titular'],
            cuenta_destino=form_data['cuenta_beneficiaria'],
            beneficiario=form_data['beneficiario'],
            importe=form_data['importe'],
            clave_autorizacion=form_data['clave_autorizacion'],
            operacion=form_data['operacion'],
            usuario=session['usuario'],
            motivo=form_data['motivo']
        )
        
        # Guardar en la base de datos
        try:
            with get_session() as session_db:
                session_db.add(nuevo_retiro)
                session_db.commit()
            flash("autorización generada con exito", "success")
        except Exception as e:
            flash(f"Error al guardar en la base de datos: {str(e)}", "danger")
            return redirect("/retiros")
        
        # Generar el documento Word
        try:
            doc = Document()
            
            # Agregar estilos y formato
            style = doc.styles['Normal']
            style.font.name = 'Geomanist'
            style.font.size = Pt(12)
            
            # Título
            title = doc.add_paragraph("Autorización de retiro")
            title.runs[0].bold = True
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Contenido
            doc.add_paragraph(f"OPERACIÓN EN {form_data['operacion']}")
            doc.add_paragraph(f"Estimado(a) compañero(a) de sucursal {form_data['cc']} {form_data['sucursal']}")
            doc.add_paragraph(
                "Con base al Front Office Retiros en Ventanilla apartado 4.82 al 4.86 y a la Circular 088 DGABC 2014 "
                "Traspasos entre Cuentas Propias y de Terceros, puede ejecutarse la siguiente operación."
            )
            
            # Detalles de la operación
            details = [
                f"Cuenta: {form_data['cuenta_origen']}",
                f"Monto: ${form_data['importe']:,.2f}",
                f"Nº de Autorización: {form_data['clave_autorizacion']}",
                f"Fecha de registro: {form_data['fecha_tramite'].strftime('%Y-%m-%d')}",
                f"Usuario atención: {session['usuario']}"
            ]
            
            for detail in details:
                p = doc.add_paragraph()
                p.add_run(detail).bold = True
            
            # Nota importante
            note = doc.add_paragraph(
                "La validación de la documentación original tanto en firmas, identificaciones oficiales, tipo de cuenta, "
                "no alteraciones, no enmendaduras, no tachaduras etc., es responsabilidad del personal de la sucursal que "
                "realice y formalice la operación y dar cumplimiento a la circular 036/DGABS/DS/2018."
            )
            note.style = doc.styles['Normal']
            
            doc.add_paragraph("Saludos")
            
            # Guardar el documento en memoria
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            
            response = make_response(send_file(
                file_stream,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f"retiro_autorizacion_{form_data['clave_autorizacion']}.docx"
            ))
            
            # Agregar headers para evitar problemas de caché
            response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
            response.headers['Pragma'] = 'no-cache'
            response.headers['Expires'] = '0'
            
            return response
            
        except Exception as e:
            flash(f"Error al generar el documento: {str(e)}", "danger")
            return redirect("/retiros")
            
    except Exception as e:
        flash(f"Error inesperado: {str(e)}", "danger")
        return redirect("/retiros")

@app.route("/concentradoretiros", methods=["GET", "POST"])
def concentradoretiro():
    if "usuario" not in session:
        return redirect("/login")

    usuario_logeado = session["usuario"]
    search_term = request.args.get("search", "")
    page = int(request.args.get("page", 1))
    per_page = 15

    with get_session() as session_db:
        # Verificar si el usuario es maestro
        user = session_db.query(usuarios).filter_by(nombre=usuario_logeado).first()
        es_maestro = user and user.es_maestro

        # Modificar la consulta según el tipo de usuario
        if es_maestro:
            query = session_db.query(Retiros)
        else:
            query = session_db.query(Retiros).filter_by(usuario=usuario_logeado)

        # Filtro de búsqueda
        if search_term:
            query = query.filter(
                Retiros.cuenta_origen.cast(String).like(f"%{search_term}%") |
                Retiros.titular.like(f"%{search_term}%") |
                Retiros.beneficiario.like(f"%{search_term}%") |
                Retiros.operacion.like(f"%{search_term}%") |
                Retiros.sucursal.like(f"%{search_term}%") |
                Retiros.usuario.like(f"%{search_term}%") |
                Retiros.cuenta_destino.like(f"%{search_term}%")
            )

        # Paginación y obtención de datos
        total = query.count()
        retiros = query.order_by(Retiros.fecha_tramite.desc()).offset((page - 1) * per_page).limit(per_page).all()

        retiros_data = [
            {
                "id": retiros.id,
                "fecha_tramite": retiros.fecha_tramite or '',
                "cc": retiros.cc or '',
                "sucursal": retiros.sucursal or '',
                "cr": retiros.cr or '',
                "cuenta_origen": retiros.cuenta_origen or '',
                "titular": retiros.titular or '',
                "cuenta_destino": retiros.cuenta_destino or '',
                "beneficiario": retiros.beneficiario or '',
                "importe": retiros.importe or '',
                "clave_autorizacion": retiros.clave_autorizacion or '',
                "operacion": retiros.operacion or '',
                "usuario": retiros.usuario or '',
                "motivo": retiros.motivo or ''
            }
            for retiros in retiros
        ]

        # Obtener el conteo de estatus para usuarios maestros o regulares
        if es_maestro:
            operacion_count = session_db.query(
                Retiros.operacion, func.count(Retiros.id).label("total")
            ).group_by(Retiros.operacion).order_by(Retiros.operacion).all()
        else:
            operacion_count = session_db.query(
                Retiros.operacion, func.count(Retiros.id).label("total")
            ).filter_by(usuario=usuario_logeado).group_by(Retiros.operacion).order_by(Retiros.operacion).all()

        # Convertir los resultados a un diccionario de totales por estatus
        operacion_totals = {row.operacion: row.total for row in operacion_count}

        # Calcular el total general a partir de los resultados de estatus_count
        total_general = sum(operacion_totals.values())

        # Agregar el total general al diccionario para facilitar su acceso
        operacion_totals["total_general"] = total_general

    return render_template(
        "concentradoretiros.html",
        retiros=retiros_data,
        page=page,
        total=total,
        per_page=per_page,
        search_term=search_term,
        operacion_totals=operacion_totals,
        total_general=total_general,
        usuario_logeado=session["usuario"],
    )

@app.route('/downloadretiros', methods=['GET', 'POST'])
def downloadretiros():
    if "usuario" not in session:
        flash("Por favor, inicia sesión para acceder a esta página.", "error")
        return redirect(url_for("login"))

    usuario_logeado = session["usuario"]

    with get_session() as session_db:
        user = session_db.query(usuarios).filter_by(nombre=usuario_logeado).first()
        es_maestro = user and user.es_maestro

        if request.method == 'POST':
            try:
                start_date = request.form.get('start_date')
                end_date = request.form.get('end_date')
                opcion_descarga = request.form.get('opcion_descarga')

                if not start_date or not end_date:
                    flash("Por favor, selecciona un rango de fechas válido.", "error")
                    return redirect(url_for("downloadretiros"))

                start_date = datetime.strptime(start_date, '%Y-%m-%d').date()
                end_date = datetime.strptime(end_date, '%Y-%m-%d').date()

                query = session_db.query(Retiros).filter(
                    Retiros.fecha_tramite >= start_date, Retiros.fecha_tramite <= end_date
                )

                if not es_maestro or (es_maestro and opcion_descarga == "propios"):
                    query = query.filter(Retiros.usuario == usuario_logeado)

                retiros = query.all()

                if not retiros:
                    flash("No se encontraron registros para la consulta.", "info")
                    return redirect(url_for("downloadretiros"))

                data = [
                    {
                        "fecha_tramite": retiros.fecha_tramite,
                        "cc": retiros.cc,
                        "sucursal": retiros.sucursal,
                        "cuenta_origen": retiros.cuenta_origen,
                        "titular": retiros.titular,
                        "cuenta_destino": retiros.cuenta_destino,
                        "beneficiario": retiros.beneficiario,
                        "importe": retiros.importe,
                        "clave_autorizacion": retiros.clave_autorizacion,
                        "operacion": retiros.operacion,
                        "usuario": retiros.usuario,
                        "motivo": retiros.motivo
                    }
                    for retiros in retiros
                ]

                df = pd.DataFrame(data)
                output = io.BytesIO()
                with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                    df.to_excel(writer, index=False, sheet_name='Retiros')

                output.seek(0)

                return send_file(
                    output,
                    as_attachment=True,
                    download_name="Concentrado_retiros.xlsx",
                    mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )

            except Exception as e:
                flash(f"Error al procesar la solicitud: {e}", "error")
                return redirect(url_for("downloadretiros"))

    return render_template("downloadretiros.html")

def generar_codigo():
    return ''.join(random.choices(string.ascii_uppercase + string.digits, k=10))

@app.route('/generar_codigo', methods=['POST'])
def generar_codigo_endpoint():
    # Aquí puedes validar los datos recibidos si es necesario
    data = request.json
    # Generar el código y la fecha actual
    codigo = generar_codigo()
    fecha_hora = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    return jsonify({'codigo': codigo, 'fecha_hora': fecha_hora})

@app.route("/concentradonominales", methods=["GET", "POST"])
def concentradonominales():
    if "usuario" not in session:
        return redirect("/login")

    usuario_logeado = session["usuario"]
    search_term = request.args.get("search", "")
    page = int(request.args.get("page", 1))
    per_page = 15

    with get_session() as session_db:
        # Verificar si el usuario es maestro
        user = session_db.query(usuarios).filter_by(nombre=usuario_logeado).first()
        es_maestro = user and user.es_maestro

        # Modificar la consulta según el tipo de usuario
        if es_maestro:
            query = session_db.query(Nominales)
        else:
            query = session_db.query(Nominales).filter_by(usuario=usuario_logeado)

        # Filtro de búsqueda
        if search_term:
            query = query.filter(
                Nominales.numero_cuenta.cast(String).like(f"%{search_term}%") |
                Nominales.beneficiario.like(f"%{search_term}%") |
                Nominales.tarjeta.like(f"%{search_term}%") |
                Nominales.cc.like(f"%{search_term}%") |
                Nominales.sucursal.like(f"%{search_term}%") |
                Nominales.usuario.like(f"%{search_term}%")
            )

        # Paginación y obtención de datos
        total = query.count()
        nominales = query.order_by(Nominales.fecha_incidencia.desc()).offset((page - 1) * per_page).limit(per_page).all()

        nominales_data = [
            {
                "id": nominales.id,
                "remesa": nominales.remesa or '',
                "numero_cuenta": nominales.numero_cuenta or '',
                "tarjeta": nominales.tarjeta or '',
                "beneficiario": nominales.beneficiario or '',
                "cc": nominales.cc or '',
                "sucursal": nominales.sucursal or '',
                "entidad": nominales.entidad or '',
                "incidencia": nominales.incidencia or '',
                "observaciones": nominales.observaciones or '',
                "fecha_incidencia": nominales.fecha_incidencia or '',
                "usuario": nominales.usuario or ''
            }
            for nominales in nominales
        ]

        # Obtener el conteo de incidencias para usuarios maestros o regulares
        if es_maestro:
            incidencia_count = session_db.query(
                Nominales.incidencia, func.count(Nominales.id).label("total")
            ).group_by(Nominales.incidencia).order_by(Nominales.incidencia).all()
        else:
            incidencia_count = session_db.query(
                Nominales.incidencia, func.count(Nominales.id).label("total")
            ).filter_by(usuario=usuario_logeado).group_by(Nominales.incidencia).order_by(Nominales.incidencia).all()

        # Convertir los resultados a un diccionario de totales por incidencias
        incidencia_totals = {row.incidencia: row.total for row in incidencia_count}

        # Calcular el total general a partir de los resultados de incidencia_count
        total_general = sum(incidencia_totals.values())

        # Agregar el total general al diccionario para facilitar su acceso
        incidencia_totals["total_general"] = total_general
    
    return render_template(
        "concentradonominales.html",
        nominales=nominales_data,
        page=page,
        total=total,
        per_page=per_page,
        search_term=search_term,
        incidencia_totals=incidencia_totals,
        total_general=total_general,
        usuario_logeado=session["usuario"],
    )

@app.route('/downloadnominales', methods=['GET', 'POST'])
def downloadnominales():
    if "usuario" not in session:
        flash("Por favor, inicia sesión para acceder a esta página.", "error")
        return redirect(url_for("login"))

    usuario_logeado = session["usuario"]

    with get_session() as session_db:
        user = session_db.query(usuarios).filter_by(nombre=usuario_logeado).first()
        es_maestro = user and user.es_maestro

        if request.method == 'POST':
            try:
                start_date = request.form.get('start_date')
                end_date = request.form.get('end_date')
                opcion_descarga = request.form.get('opcion_descarga')

                if not start_date or not end_date:
                    flash("Por favor, selecciona un rango de fechas válido.", "error")
                    return redirect(url_for("downloadnominales"))

                start_date = datetime.strptime(start_date, '%Y-%m-%d').date()
                end_date = datetime.strptime(end_date, '%Y-%m-%d').date()

                query = session_db.query(Nominales).filter(
                    Nominales.fecha_incidencia >= start_date, Nominales.fecha_incidencia <= end_date
                )

                if not es_maestro or (es_maestro and opcion_descarga == "propios"):
                    query = query.filter(Nominales.usuario == usuario_logeado)

                nominales = query.all()

                if not nominales:
                    flash("No se encontraron registros para la consulta.", "info")
                    return redirect(url_for("downloadnominales"))

                data = [
                    {
                        "REMESA": nominales.remesa,
                        "NUMERO DE CUENTA": nominales.numero_cuenta,
                        "NUMERO DE TARJETA": nominales.tarjeta,
                        "BENEFICIARIO": nominales.beneficiario,
                        "CC": nominales.cc,
                        "SUCURSAL": nominales.sucursal,
                        "ENTIDAD": nominales.entidad,
                        "INCIDENCIA": nominales.incidencia,
                        "OBSERVACIONES": nominales.observaciones,
                        "FECHA DE INCIDENCIA": nominales.fecha_incidencia,
                        "USUARIO": nominales.usuario
                    }
                    for nominales in nominales
                ]

                df = pd.DataFrame(data)
                output = io.BytesIO()
                with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                    df.to_excel(writer, index=False, sheet_name='Nominales')

                output.seek(0)

                return send_file(
                    output,
                    as_attachment=True,
                    download_name="Concentrado_Nominales.xlsx",
                    mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )

            except Exception as e:
                flash(f"Error al procesar la solicitud: {e}", "error")
                return redirect(url_for("downloadnominales"))

    return render_template("downloadnominales.html")

@app.route('/uploadnominales', methods=['GET', 'POST'])
def upload_nomiles():
    if "usuario" not in session:
        return redirect("/login")

    if request.method == 'GET':
        return render_template('uploadnominales.html')

    try:
        if 'file' not in request.files:
            flash('No se seleccionó ningún archivo.', 'error')
            return redirect('/uploadnominales')

        file = request.files['file']
        if file.filename == '':
            flash('El nombre del archivo está vacío.', 'error')
            return redirect(url_for('uploadnominales'))

        df = pd.read_excel(file)
        required_columns = ["REMESA", "NUMERO DE CUENTA", "TARJETA", "BENEFICIARIO", "CC", "SUCURSAL", "ENTIDAD", "INCIDENCIA"]
        for column in required_columns:
            if column not in df.columns:
                flash(f"El archivo no contiene la columna obligatoria: {column}", 'error')
                return redirect('/uploadnominales')

        # Limpiar caracteres en columnas de texto
        text_columns = ["REMESA", "BENEFICIARIO", "SUCURSAL", "ENTIDAD", "INCIDENCIA", "OBSERVACIONES"]
        for column in text_columns:
            if column in df.columns:
                df[column] = df[column].apply(clean_sql_string)

        df = df.dropna(subset=required_columns)
        if df.empty:
            flash("No hay filas válidas en el archivo después de filtrar valores nulos.", 'error')
            return redirect('/uploadnominales')

        # Reemplazar valores NaN con cadenas vacías en todo el DataFrame
        df = df.fillna("")

        numeric_columns = [
            "NUMERO DE CUENTA",
            "CC",
        ]
        for column in numeric_columns:
            if column in df.columns:
                # Primero limpiar cualquier carácter no numérico
                df[column] = df[column].astype(str).apply(lambda x: ''.join(filter(str.isdigit, str(x))))
                df[column] = pd.to_numeric(df[column], errors='coerce').fillna(0).astype(int)

        fecha_incidencia = date.today()
        duplicados = []
        errores_cc = []
        invalid_rows = []
        caracteres_especiales = []

        with get_session() as session_db:
            for _, row in df.iterrows():
                sucursal = clean_sql_string(row["SUCURSAL"])
                numero_cuenta = row["NUMERO DE CUENTA"]
                cc = row["CC"]
                entidad = row["ENTIDAD"]

                # Verificar si hubo limpieza de caracteres especiales
                original_beneficiario = str(row["BENEFICIARIO"])
                cleaned_beneficiario = clean_sql_string(original_beneficiario)
                if original_beneficiario != cleaned_beneficiario:
                    caracteres_especiales.append(f"Cuenta: {numero_cuenta}")

                if len(str(numero_cuenta)) < 8 or len(str(numero_cuenta)) > 10:
                    invalid_rows.append(row.to_dict())
                    continue

                if cc == 0:
                    errores_cc.append(f"Cuenta: {numero_cuenta} con CC igual a 0.")
                    continue

                if session_db.query(Nominales).filter_by(numero_cuenta=numero_cuenta).first():
                    duplicados.append(numero_cuenta)
                    continue

                if pd.isna(sucursal) or sucursal == "":
                    sucursal = "SUCURSAL NO CARGADA CORRECTAMENTE"
                    continue

                nominales = Nominales(
                    remesa=row.get("REMESA"),
                    numero_cuenta=numero_cuenta,
                    tarjeta=row.get("TARJETA"),
                    beneficiario=cleaned_beneficiario,
                    cc=cc,
                    sucursal=sucursal,
                    entidad=entidad,
                    incidencia=clean_sql_string(row.get("INCIDENCIA", "")),
                    observaciones=clean_sql_string(row.get("OBSERVACIONES", "")),
                    fecha_incidencia=fecha_incidencia,
                    usuario=session["usuario"],
                )
                session_db.add(nominales)

            session_db.commit()

        if caracteres_especiales:
            flash(f"Se limpiaron caracteres especiales en {len(caracteres_especiales)} registros", 'info')

        if duplicados:
            flash(f"Algunas cuentas ya existían y no se cargaron: {', '.join(map(str, duplicados))}", 'warning')

        if errores_cc:
            flash(f"Errores de CC (igual a 0): {', '.join(errores_cc)}", 'error')

        if invalid_rows:
            flash(f"Filas inválidas (número de cuenta no tiene entre 8 y 10 dígitos): {len(invalid_rows)}", 'warning')

        flash("Archivo procesado y datos cargados en la base de datos con éxito.", 'success')
    except Exception as e:
        flash(f"Error al procesar el archivo: {str(e)}", 'error')

    return redirect('/uploadnominales')

if __name__ == '__main__':
    app.run(debug=True)